"""
Simple, unified DOCX creator - ONE function to rule them all!
No external converters, no markdown parsing, direct DOCX creation.
"""

import os
import io
import base64
import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_chapter_content(doc, chapter_data, chapter_name, is_main_title=False):
    """
    Add a chapter's content to an existing document.
    
    Args:
        doc: Document object to add content to
        chapter_data: Dictionary with chapter analysis data
        chapter_name: Name of the chapter
        is_main_title: If True, adds as main document title, else as chapter heading
    """
    if is_main_title:
        # Add as main document title
        title = doc.add_paragraph()
        title.style = doc.styles['Custom Title']
        clean_title = _format_chapter_title(chapter_name, chapter_data)
        title.add_run(clean_title)
    else:
        # Add as chapter heading (for combined documents)
        chapter_heading = doc.add_paragraph()
        chapter_heading.style = doc.styles['Custom Heading']
        clean_title = _format_chapter_title(chapter_name, chapter_data)
        chapter_heading.add_run(clean_title)
    
    # Add mindmap section - use consistent mindmaps structure
    mindmaps = chapter_data.get('mindmaps', {})
    comprehensive_mindmap = mindmaps.get('comprehensive_mindmap') if mindmaps else None
    
    if comprehensive_mindmap:
        _add_mindmap_section(doc, comprehensive_mindmap)
    else:
        # Check if mindmap is stored at top level (wrong structure)
        if 'comprehensive_mindmap' in chapter_data:
            _add_mindmap_section(doc, chapter_data['comprehensive_mindmap'])
    
    # Add explanation section  
    if chapter_data.get('mindmap_explanation'):
        _add_explanation_section(doc, chapter_data['mindmap_explanation'])
    
    # Add summary section
    if chapter_data.get('quick_summary'):
        # Add simple summary heading
        summary_heading = doc.add_paragraph()
        summary_heading.style = doc.styles['Custom SubHeading']
        summary_heading.add_run("Summary")
        _add_summary_section(doc, chapter_data['quick_summary'])
    
    # Add detailed analysis (was missing in combined docs!)
    if chapter_data.get('analysis_summary'):
        _add_analysis_section(doc, chapter_data['analysis_summary'])


def create_docx(chapters, output_path=None):
    """
    UNIFIED DOCX creation function - handles both single and multiple chapters!
    
    Args:
        chapters: Either:
                 - Single chapter: (chapter_name, chapter_data) tuple
                 - Multiple chapters: dict of {chapter_name: chapter_data}
        output_path: Where to save (if None, returns BytesIO)
    
    Returns:
        Path to saved file or BytesIO object
    """
    # Create document
    doc = Document()
    _setup_styles(doc)
    
    # Normalize input to handle both single and multiple chapters
    if isinstance(chapters, tuple) and len(chapters) == 2:
        # Single chapter: (chapter_name, chapter_data)
        chapter_name, chapter_data = chapters
        chapters_dict = {chapter_name: chapter_data}
        is_single_chapter = True
    elif isinstance(chapters, dict):
        # Multiple chapters: {chapter_name: chapter_data}
        chapters_dict = chapters
        is_single_chapter = len(chapters_dict) == 1
    else:
        raise ValueError("chapters must be either (chapter_name, chapter_data) tuple or {chapter_name: chapter_data} dict")
    
    if is_single_chapter:
        # Single chapter document
        chapter_name, chapter_data = next(iter(chapters_dict.items()))
        _add_chapter_content(doc, chapter_data, chapter_name, is_main_title=True)
    else:
        # Multi-chapter document
        # Add main document title
        title = doc.add_paragraph()
        title.style = doc.styles['Custom Title']
        title.add_run("Mind Map Analysis Report")
        
        # Add each chapter
        for i, (chapter_name, chapter_data) in enumerate(chapters_dict.items()):
            _add_chapter_content(doc, chapter_data, chapter_name, is_main_title=False)
            
            # Add page break between chapters (but not after the last one)
            if i < len(chapters_dict) - 1:
                doc.add_page_break()
    
    # Save or return
    if output_path:
        doc.save(output_path)
        return output_path
    else:
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io


# Legacy compatibility functions
def create_chapter_docx_direct(chapter_data, chapter_name, output_path=None):
    """Legacy compatibility wrapper for single chapter creation"""
    return create_docx((chapter_name, chapter_data), output_path)


def _format_chapter_title(chapter_name, chapter_data=None):
    """
    Convert raw chapter filename to a professional, readable title.
    
    Examples:
        "07_chapter_chapter-2-how-to-almost-make-anything" -> "Chapter 2: How to Almost Make Anything"
        "03_introduction_getting-started" -> "Chapter 3: Introduction - Getting Started"
    """
    # First try to get title from chapter_data if available
    if chapter_data:
        if 'chapter_title' in chapter_data:
            return chapter_data['chapter_title']
        if 'title' in chapter_data:
            return chapter_data['title']
    
    # Extract chapter number and title from filename
    import re
    
    # Remove file extension if present
    clean_name = chapter_name.replace('.md', '').replace('.docx', '')
    
    # Pattern to match: optional number + chapter + actual title
    # Handles formats like: "07_chapter_chapter-2-how-to-almost-make-anything"
    pattern = r'^(\d+)_?chapter_?(.*)$'
    match = re.match(pattern, clean_name, re.IGNORECASE)
    
    if match:
        chapter_num = match.group(1)
        title_part = match.group(2)
        
        # Clean up the title part
        title_part = title_part.replace('chapter-', '').replace('chapter_', '')
        title_part = title_part.replace('-', ' ').replace('_', ' ')
        title_part = title_part.strip()
        
        # Extract chapter number from title if it exists (e.g., "2-how-to-make")
        title_num_pattern = r'^(\d+)[-_\s]*(.*)'
        title_match = re.match(title_num_pattern, title_part)
        if title_match:
            actual_chapter_num = title_match.group(1)
            actual_title = title_match.group(2)
        else:
            actual_chapter_num = chapter_num
            actual_title = title_part
        
        # Capitalize title properly
        actual_title = actual_title.replace('-', ' ').replace('_', ' ')
        actual_title = ' '.join(word.capitalize() for word in actual_title.split())
        
        return f"Chapter {actual_chapter_num}: {actual_title}"
    
    # Fallback: just clean up the name
    clean_title = clean_name.replace('-', ' ').replace('_', ' ')
    clean_title = ' '.join(word.capitalize() for word in clean_title.split())
    return clean_title


def _clean_markdown_remnants(content):
    """
    Remove markdown formatting remnants from content.
    This handles inline formatting that shouldn't appear in Word docs.
    """
    import re
    
    if not content:
        return content
    
    # Remove horizontal rules
    content = re.sub(r'^[-=_*]{3,}$', '', content, flags=re.MULTILINE)
    
    # Clean up bold formatting that's not at line boundaries
    # Replace **text** with just text (will be handled by inline formatting)
    content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
    
    # Clean up italic formatting
    content = re.sub(r'\*(.*?)\*', r'\1', content)
    content = re.sub(r'_(.*?)_', r'\1', content)
    
    # Remove extra markdown syntax
    content = re.sub(r'`([^`]+)`', r'\1', content)  # Remove code backticks
    content = re.sub(r'~~(.*?)~~', r'\1', content)  # Remove strikethrough
    
    # Clean up multiple consecutive spaces/newlines
    content = re.sub(r'\n{3,}', '\n\n', content)  # Max 2 consecutive newlines
    content = re.sub(r' {2,}', ' ', content)  # Multiple spaces to single space
    
    return content.strip()


def _add_paragraph_with_inline_formatting(paragraph, text):
    """
    Add text to a paragraph while preserving bold formatting from original markdown.
    This function specifically looks for **bold** patterns and applies Word formatting.
    """
    import re
    
    # Find all **bold** patterns in the text
    bold_pattern = r'\*\*(.*?)\*\*'
    parts = re.split(bold_pattern, text)
    
    for i, part in enumerate(parts):
        if not part:  # Skip empty strings
            continue
            
        if i % 2 == 1:  # Odd indices are the bold text (captured groups)
            run = paragraph.add_run(part)
            run.bold = True
        else:  # Even indices are regular text
            paragraph.add_run(part)


def _setup_styles(doc):
    """Setup document styles with proper Calibri font"""
    styles = doc.styles
    
    # Title style
    if 'Custom Title' not in [s.name for s in styles]:
        title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Calibri'
        title_font.size = Pt(22)
        title_font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(18)
        title_style.paragraph_format.space_before = Pt(12)

    # Heading style
    if 'Custom Heading' not in [s.name for s in styles]:
        heading_style = styles.add_style('Custom Heading', WD_STYLE_TYPE.PARAGRAPH)
        heading_font = heading_style.font
        heading_font.name = 'Calibri'
        heading_font.size = Pt(16)
        heading_font.bold = True
        heading_font.color.rgb = None  # Default color
        heading_style.paragraph_format.space_before = Pt(14)
        heading_style.paragraph_format.space_after = Pt(8)

    # Body style
    if 'Custom Body' not in [s.name for s in styles]:
        body_style = styles.add_style('Custom Body', WD_STYLE_TYPE.PARAGRAPH)
        body_font = body_style.font
        body_font.name = 'Calibri'
        body_font.size = Pt(11)
        body_style.paragraph_format.space_after = Pt(6)
        body_style.paragraph_format.line_spacing = 1.15
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
    # Sub-heading style
    if 'Custom SubHeading' not in [s.name for s in styles]:
        subheading_style = styles.add_style('Custom SubHeading', WD_STYLE_TYPE.PARAGRAPH)
        subheading_font = subheading_style.font
        subheading_font.name = 'Calibri'
        subheading_font.size = Pt(13)
        subheading_font.bold = True
        subheading_style.paragraph_format.space_before = Pt(10)
        subheading_style.paragraph_format.space_after = Pt(6)


def _add_mindmap_section(doc, mindmap_content):
    """Add mindmap section with image or text fallback"""
    # Try to generate and add mindmap image
    image_added = False
    try:
        # Clean the mermaid content
        cleaned_content = _clean_mermaid_for_api(mindmap_content)
        
        # Generate image
        image_data = _generate_mindmap_image(cleaned_content)
        
        if image_data:
            # Add image to document using the simple approach
            image_stream = io.BytesIO(image_data)
            
            # Create centered paragraph for image
            image_paragraph = doc.add_paragraph()
            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add picture directly to the paragraph
            run = image_paragraph.add_run()
            run.add_picture(image_stream, width=Inches(6))
            
            print("✓ Mindmap image successfully added to DOCX")
            image_added = True
            
    except Exception as e:
        print(f"✗ Error adding mindmap image: {e}")
        print(f"Mindmap content preview: {str(mindmap_content)[:100]}...")
    
    # Add text fallback if image failed
    if not image_added:
        print("📝 Adding mindmap text fallback to DOCX")
        
        # Add a note about the visualization
        note_para = doc.add_paragraph()
        note_para.style = doc.styles['Custom Body']
        note_para.add_run("Note: Mindmap visualization not available. The mindmap structure is shown below:")
        
        # Add the mindmap content as code block
        code_para = doc.add_paragraph()
        code_para.style = doc.styles['Custom Body']
        code_run = code_para.add_run(str(mindmap_content))
        code_run.font.name = 'Consolas'
        code_run.font.size = Pt(9)


def _add_explanation_section(doc, explanation_content):
    """Add explanation section with proper markdown formatting"""
    # Clean content to remove duplicate headers
    cleaned_content = _clean_content_headers(explanation_content, 
                                            ["Mind Map Explanation", "Mind Map Notes"])
    
    # Content already has its own meaningful title, no need for generic header
    
    _add_formatted_content(doc, cleaned_content)


def _add_summary_section(doc, summary_content):
    """Add summary section with proper markdown formatting"""
    # Clean content to remove duplicate headers
    cleaned_content = _clean_content_headers(summary_content, 
                                            ["Comprehensive Summary", "Quick Summary", "Summary"])
    
    # Content already has its own meaningful title, no need for generic header
    
    _add_formatted_content(doc, cleaned_content)


def _add_analysis_section(doc, analysis_content):
    """Add detailed analysis section with proper markdown formatting"""
    # Content already has its own meaningful title, no need for generic header
    
    _add_formatted_content(doc, analysis_content)


def _add_formatted_content(doc, content):
    """Convert markdown-style content to properly formatted Word content"""
    if not content:
        return
        
    content_str = str(content)
    # Clean all markdown remnants BEFORE processing
    content_str = _clean_markdown_remnants(content_str)
    lines = content_str.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            # Add small spacing for empty lines
            doc.add_paragraph()
            continue
            
        # Skip horizontal rules completely
        if line in ['---', '===', '***', '___']:
            continue
            
        # Handle different markdown elements
        if line.startswith('###'):
            # Sub-sub heading
            para = doc.add_paragraph()
            para.style = doc.styles['Custom SubHeading']
            run = para.add_run(line.replace('###', '').strip())
            run.font.size = Pt(11)
            
        elif line.startswith('##'):
            # Sub heading  
            para = doc.add_paragraph()
            para.style = doc.styles['Custom SubHeading']
            run = para.add_run(line.replace('##', '').strip())
            run.font.size = Pt(12)
            
        elif line.startswith('#'):
            # Main heading
            para = doc.add_paragraph()
            para.style = doc.styles['Custom Heading']
            para.add_run(line.replace('#', '').strip())
            
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet point
            para = doc.add_paragraph()
            para.style = doc.styles['Custom Body']
            para.style.paragraph_format.left_indent = Inches(0.25)
            para.add_run('• ' + line[2:].strip())
            
        elif line.startswith('**') and line.endswith('**') and line.count('**') == 2:
            # Bold text (entire line)
            para = doc.add_paragraph()
            para.style = doc.styles['Custom Body']
            run = para.add_run(line.replace('**', ''))
            run.bold = True
            
        elif ':' in line and len(line.split(':')) == 2:
            # Key-value pairs (often in summaries)
            para = doc.add_paragraph()
            para.style = doc.styles['Custom Body']
            key, value = line.split(':', 1)
            key_run = para.add_run(key.strip() + ': ')
            key_run.bold = True
            para.add_run(value.strip())
            
        else:
            # Regular paragraph with inline formatting
            para = doc.add_paragraph()
            para.style = doc.styles['Custom Body']
            _add_paragraph_with_inline_formatting(para, line)


def _add_text_fallback(doc, content, title):
    """Add text content as fallback"""
    heading = doc.add_paragraph()
    heading.style = doc.styles['Custom Heading']
    heading.add_run(f"{title} (Text Version)")
    
    para = doc.add_paragraph()
    para.style = doc.styles['Custom Body']
    para.add_run(str(content))


def _clean_mermaid_for_api(content):
    """Clean mermaid content for API compatibility"""
    if not content:
        return ""
    
    # The API actually DOES accept 'mindmap' - don't replace it
    # Just clean up any problematic formatting
    cleaned = str(content).strip()
    
    # Remove any existing markdown code blocks if present
    if cleaned.startswith('```mermaid'):
        cleaned = cleaned.replace('```mermaid\n', '').replace('```', '')
    
    # Make sure it starts with a valid diagram type
    if not cleaned.startswith('mindmap'):
        # If it doesn't start with mindmap, try to fix it
        if 'mindmap' in cleaned:
            # Extract from the first occurrence of mindmap
            mindmap_index = cleaned.find('mindmap')
            cleaned = cleaned[mindmap_index:]
        elif 'visualmap' in cleaned:
            # Replace visualmap back to mindmap
            cleaned = cleaned.replace('visualmap', 'mindmap')
    
    # Content is now generated with correct syntax from the LLM
    # Just validate that it starts with mindmap
    if not cleaned.startswith('mindmap'):
        print("⚠ Content doesn't appear to be a valid mindmap")
        return ""
    
    return cleaned


def _clean_content_headers(content, headers_to_remove):
    """
    Remove duplicate headers from content that conflict with DOCX section headers
    
    Args:
        content: Text content that may contain duplicate headers
        headers_to_remove: List of header text to remove (e.g., ["Mind Map Explanation"])
        
    Returns:
        Cleaned content without duplicate headers
    """
    if not content:
        return content
    
    lines = content.split('\n')
    cleaned_lines = []
    skip_next_separator = False
    
    for line in lines:
        line_stripped = line.strip()
        
        # Skip horizontal separators after removed headers
        if skip_next_separator and line_stripped in ['---', '===', '***']:
            skip_next_separator = False
            continue
        
        # Check if this line contains any of the headers to remove
        should_skip = False
        for header_text in headers_to_remove:
            if line_stripped.startswith('#') and header_text.lower() in line_stripped.lower():
                should_skip = True
                skip_next_separator = True
                break
        
        if not should_skip:
            cleaned_lines.append(line)
            skip_next_separator = False
    
    return '\n'.join(cleaned_lines)


def _generate_mindmap_image(mermaid_content):
    """Generate mindmap image using Mermaid API with robust fallbacks"""
    try:
        print(f"🔄 Generating mindmap image...")
        
        # Ensure we have valid mermaid content
        if not mermaid_content or len(str(mermaid_content).strip()) < 10:
            print("⚠ Mindmap content too short or empty")
            return None
            
        # Clean and prepare the mermaid diagram
        cleaned_content = str(mermaid_content).strip()
        
        # Remove any existing markdown code blocks (common source of 400 errors)
        if cleaned_content.startswith('```mermaid'):
            cleaned_content = cleaned_content.replace('```mermaid\n', '').replace('```', '')
        elif cleaned_content.startswith('```'):
            # Remove any generic code blocks
            lines = cleaned_content.split('\n')
            if lines[0].strip() == '```' or lines[0].startswith('```'):
                lines = lines[1:]
            if lines and lines[-1].strip() == '```':
                lines = lines[:-1]
            cleaned_content = '\n'.join(lines)
        
        # Fix visualmap back to mindmap (our cleaning might have broken it)
        if cleaned_content.startswith('visualmap'):
            cleaned_content = cleaned_content.replace('visualmap', 'mindmap', 1)
        
        # Ensure it starts with mindmap
        if not cleaned_content.startswith('mindmap'):
            if 'mindmap' in cleaned_content.lower():
                # Try to extract the mindmap portion
                mindmap_index = cleaned_content.lower().find('mindmap')
                cleaned_content = cleaned_content[mindmap_index:]
            else:
                print("⚠ Content doesn't appear to be a valid mindmap")
                return None
        
        # Validate mindmap syntax
        if not _validate_mindmap_syntax(cleaned_content):
            print("⚠ Invalid mindmap syntax detected")
            return None
        
    
        
        # CRITICAL FIX: mermaid.ink expects RAW mermaid code, NOT wrapped in markdown
        # Encode the cleaned content directly
        encoded = base64.b64encode(cleaned_content.encode('utf-8')).decode('utf-8')
        
        # Try multiple API endpoints for robustness
        api_attempts = [
            {"url": f"https://mermaid.ink/img/{encoded}", "format": "PNG"},
            {"url": f"https://mermaid.ink/img/{encoded}?type=png", "format": "PNG"},
            {"url": f"https://mermaid.ink/img/{encoded}?bgColor=white", "format": "PNG with white background"},
        ]
        
        for i, attempt in enumerate(api_attempts, 1):
            try:
                print(f"🌐 Attempt {i}: Calling Mermaid API ({attempt['format']})...")
                print(f"   URL: {attempt['url'][:70]}...")
                
                response = requests.get(attempt['url'], timeout=30)
                
                if response.status_code == 200:
                    print(f"✓ Mindmap image generated successfully ({len(response.content)} bytes)")
                    return response.content
                else:
                    print(f"✗ Attempt {i} failed: {response.status_code} - {response.text[:100]}")
                    if i == 1:  # Only show encoding details on first failure
                        print(f"📝 Encoded content (first 200 chars): {encoded[:200]}...")
                        
            except requests.exceptions.RequestException as e:
                print(f"✗ Network error on attempt {i}: {e}")
                continue
        
        print("⚠ All API attempts failed")
        return None
            
    except Exception as e:
        print(f"✗ Error generating mindmap image: {e}")
        import traceback
        traceback.print_exc()
        return None


def _validate_mindmap_syntax(content):
    """Basic validation of mindmap syntax"""
    try:
        lines = content.strip().split('\n')
        if not lines:
            return False
        
        # Must start with 'mindmap'
        if not lines[0].strip().startswith('mindmap'):
            return False
        
        # Should have at least a root node with correct syntax root((text))
        has_valid_root = False
        for line in lines[1:5]:  # Check first few lines after 'mindmap'
            stripped = line.strip()
            if stripped.startswith('root((') and stripped.endswith('))'):
                has_valid_root = True
                break
        
        return has_valid_root
        
    except Exception:
        return False


def create_combined_docx_direct(all_chapters_data, output_path=None):
    """Legacy compatibility wrapper for combined chapter creation"""
    return create_docx(all_chapters_data, output_path)
